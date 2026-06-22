#!/usr/bin/env python3
"""Genera documentacion final PDF: Times New Roman 12, interlineado 1.5."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, Inches
from docx.oxml.ns import qn

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak,
)

BASE = Path(__file__).resolve().parent
OUT_DOCX = BASE / "Documentacion_Final_GreenSQA.docx"
OUT_PDF = BASE / "Documentacion_Final_GreenSQA.pdf"

FONT = "Times-Roman"
FONT_BOLD = "Times-Bold"
SIZE = 12
LEADING = 18  # 1.5 x 12pt


def set_run_font(run, size=SIZE, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold


def build_docx():
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.18)
        section.right_margin = Inches(1.18)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(SIZE)
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def add_p(text, bold=False):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = p.add_run(text)
        set_run_font(run, bold=bold)

    def add_h(text, size=14):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run(text)
        set_run_font(run, size=size, bold=True)

    add_h("DOCUMENTACION FINAL", 16)
    add_h("Prueba Tecnica QA Automatizador — GreenSQA", 14)
    add_p("Fecha de ejecucion: 22 de junio de 2026")
    add_p("Repositorio: https://github.com/Angelicablandonn/prueba-atm-greensqa")

    sections = get_content_sections()
    for sec in sections:
        add_h(sec["title"])
        for block in sec["blocks"]:
            if block["type"] == "text":
                add_p(block["text"], bold=block.get("bold", False))
            elif block["type"] == "bullets":
                for b in block["items"]:
                    add_p(f"• {b}")

    doc.save(OUT_DOCX)


def get_content_sections():
    return [
        {
            "title": "1. Resumen Ejecutivo",
            "blocks": [
                {"type": "text", "text": (
                    "Se valido el entorno, se ejecuto el data-generator con exito (tests, JAR, CSV valido), "
                    "se corrigieron defectos del framework UI (Serenity/Cucumber) y se ejecutaron tres escenarios "
                    "K6 contra https://test.k6.io sin errores HTTP."
                )},
                {"type": "text", "text": (
                    "La automatizacion UI contra LATAM no completo escenarios por proteccion anti-bot Akamai "
                    "(Access Denied en headless; timeouts en headed). El framework compila, ejecuta y genera Serenity."
                )},
            ],
        },
        {
            "title": "2. Entorno Utilizado",
            "blocks": [{"type": "text", "text": (
                "JDK 17.0.12 | Maven 3.9.12 | Chrome 149.0.7827.115 | ChromeDriver WebDriverManager | K6 v0.57.0"
            )}],
        },
        {
            "title": "3. Resultados Data Generator",
            "blocks": [
                {"type": "text", "text": "mvn clean test: BUILD SUCCESS — 8 tests, 0 fallos."},
                {"type": "text", "text": "mvn clean package: JAR data-generator.jar (~17 MB)."},
                {"type": "text", "text": "Serial 100 registros OK (242 ms). Paralelo 500 registros OK (237 ms, 4 hilos)."},
                {"type": "text", "text": "Validacion CSV: VALIDO — 500 filas, 0 duplicados documento y nombre+apellido."},
            ],
        },
        {
            "title": "4. Resultados Automatizacion UI",
            "blocks": [
                {"type": "text", "text": "Headless (mvn verify -Dheadless.mode=true): 0 pass, 3 errors — Access Denied Akamai."},
                {"type": "text", "text": "Headed (mvn verify -Dheadless.mode=false): 0 pass, 1 failure, 2 errors — timeout/origen."},
                {"type": "text", "text": "CSV integrado: data-generator/output/datos_actual.csv (500 registros)."},
                {"type": "text", "text": "Reporte Serenity: automation/target/site/serenity/index.html"},
            ],
        },
        {
            "title": "5. Resultados K6",
            "blocks": [
                {"type": "text", "text": "Target: https://test.k6.io (no produccion)."},
                {"type": "text", "text": "prueba-pequena-10hilos.js: ~31s, 0% errores, checks 100%."},
                {"type": "text", "text": "escenario1-carga.js: ~5 min, 0% errores, checks 100%."},
                {"type": "text", "text": "escenario2-capacidad.js: ~6 min, 0% errores, checks 100%."},
            ],
        },
        {
            "title": "6. Correcciones Aplicadas",
            "blocks": [{"type": "bullets", "items": [
                "Paso Cucumber duplicado eliminado.",
                "API Serenity 4.x en Questions.",
                "Dependencia cucumber-junit-platform-engine.",
                "Selectores ampliados en HomePage.java.",
                "Anti-deteccion Chrome y pageLoadStrategy=eager.",
                "Deteccion Access Denied en AbrirHomeLatam.",
                "Ejecucion secuencial (parallel=none).",
                "Compilacion JDK 17 con -Dmaven.compiler.release=17.",
            ]}],
        },
        {
            "title": "7. Riesgos y Pendientes",
            "blocks": [
                {"type": "text", "text": "Riesgo principal: anti-bot LATAM bloquea automatizacion UI."},
                {"type": "bullets", "items": [
                    "Ejecutar UI desde red no bloqueada.",
                    "Instalar JDK 21 si es requisito.",
                    "Validar selectores cuando el sitio permita acceso.",
                ]},
            ],
        },
        {
            "title": "8. Evidencias",
            "blocks": [{"type": "bullets", "items": [
                "evidencias/data-generator-*.txt",
                "evidencias/automation-verify.txt",
                "evidencias/k6-*.txt y k6-*.json",
                "evidencias/serenity/*.png",
                "evidencias/RESUMEN.md",
            ]}],
        },
        {
            "title": "9. Decisiones Tecnicas",
            "blocks": [{"type": "bullets", "items": [
                "D1: JDK 17 por disponibilidad en entorno.",
                "D2: Deteccion temprana Access Denied.",
                "D3: No clic en Ida y vuelta (default).",
                "D4-D6: Config Chrome y pageLoadStrategy correcto.",
                "D7: K6 binario local.",
                "D8: No simular PASS — evidencia auditable.",
            ]}],
        },
    ]


def build_pdf():
    doc = SimpleDocTemplate(
        str(OUT_PDF), pagesize=letter,
        leftMargin=inch, rightMargin=inch,
        topMargin=inch, bottomMargin=inch,
    )
    title_s = ParagraphStyle("title", fontName=FONT_BOLD, fontSize=16, leading=24,
                               alignment=1, spaceAfter=12)
    h2 = ParagraphStyle("h2", fontName=FONT_BOLD, fontSize=14, leading=21,
                        spaceBefore=14, spaceAfter=8)
    body = ParagraphStyle("body", fontName=FONT, fontSize=SIZE, leading=LEADING, spaceAfter=6)
    body_b = ParagraphStyle("body_b", parent=body, fontName=FONT_BOLD)

    story = [
        Paragraph("DOCUMENTACION FINAL", title_s),
        Paragraph("Prueba Tecnica QA Automatizador — GreenSQA", h2),
        Paragraph("Fecha: 22 de junio de 2026", body),
        Paragraph("Repositorio: github.com/Angelicablandonn/prueba-atm-greensqa", body),
        Spacer(1, 12),
    ]

    for sec in get_content_sections():
        story.append(Paragraph(sec["title"], h2))
        for block in sec["blocks"]:
            if block["type"] == "text":
                style = body_b if block.get("bold") else body
                story.append(Paragraph(block["text"], style))
            elif block["type"] == "bullets":
                for item in block["items"]:
                    story.append(Paragraph(f"• {item}", body))

    # Tabla resumen evidencias
    story.append(Spacer(1, 12))
    story.append(Paragraph("Tabla de Evidencias", h2))
    data = [
        ["Prueba", "Estado", "Archivo"],
        ["Data Generator Tests", "PASO", "data-generator-tests.txt"],
        ["Data Generator Package", "PASO", "data-generator-package.txt"],
        ["CSV Validation", "PASO", "data-generator-validation.txt"],
        ["Automation Headless", "FALLO", "automation-verify.txt"],
        ["K6 Small/Load/Capacity", "PASO", "k6-*.txt / k6-*.json"],
        ["Serenity Report", "GENERADO", "serenity/*.png"],
    ]
    t = Table(data, colWidths=[2.2 * inch, 1.0 * inch, 2.3 * inch])
    t.setStyle(TableStyle([
        ("FONTNAME", (0, 0), (-1, 0), FONT_BOLD),
        ("FONTNAME", (0, 1), (-1, -1), FONT),
        ("FONTSIZE", (0, 0), (-1, -1), SIZE),
        ("LEADING", (0, 0), (-1, -1), LEADING),
        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.black),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.whitesmoke]),
    ]))
    story.append(t)

    doc.build(story)
    print(f"PDF generado: {OUT_PDF} ({OUT_PDF.stat().st_size} bytes)")


if __name__ == "__main__":
    build_docx()
    print(f"DOCX generado: {OUT_DOCX}")
    build_pdf()
